from docx import Document
from docx.shared import Inches

import pytz
from io import BytesIO

from pydicom.dataset import Dataset
from pydicom.sequence import Sequence

from pydicom.filewriter import multi_string

from pydicom.uid import UID as uid
from pydicom.filewriter import multi_string
import os
import logging
import time


from datetime import datetime

def generate_sr(patient_name="Шмидхубер Юрген Иванович",
                  doctor_name="Пупкин В.П.",
                  date=str(datetime.now()),
                  description="",
                  pathologies=[]):
    """
    Generates DICOM SR based on request data
    """
    MosmedReport(patient_name,doctor_name,date,description,pathologies).save(".")

def generate_docx(patient_name="Шмидхубер Юрген Иванович",
                  doctor_name="Пупкин В.П.",
                  date=str(datetime.now()),
                  description="",
                  pathologies=[]
                 ):
    """
    Generates printable DOCX report based on request data
    """
    document = Document()

    document.add_heading('ПРОТОКОЛ', 0)
    document.add_heading('Рентгенография легких', 1)

    p0 = document.add_paragraph(f"""
    ФИО: {patient_name}
    Дата рождения: 09.03.1997г.; Пол: М""")
    p1 = document.add_paragraph('0,04 мЗв РГ органов грудной клетки в прямой проекции в ортопозиции.')
    p2 = document.add_paragraph(description)
    document.add_heading("Заключение:",1)
    if len(pathologies)==0:
        p3 = document.add_paragraph("""
        Рентгенологические признаки очаговой патологии органов грудной клетки не выявлены.""")
    else:
        p3 = document.add_paragraph("""
        Рентгенологические признаки очаговой патологии органов грудной клетки обнаружены, необходима консультация специалиста.""")
    p5 = document.add_paragraph(f"""
    Врач: {doctor_name} Дата: {date} 8:31 Подпись:__________""")

    document.save('demo.docx')

def generate_UID(original_UID,add_id,model_id=1027):
    return uid(str(original_UID)+"."+str(model_id)+"."+str(add_id))

def get_formatted_string(value):
    """Formats a dicom header string properly
    According to http://dicom.nema.org/dicom/2013/output/chtml/part05/sect_6.2.html
    strings need to be an even number of characters. If necessary, pad with a
    single space.
    """
    formatted_string = multi_string(value)
    if len(formatted_string) % 2 != 0:
        formatted_string = formatted_string + ' '
    return formatted_string



class MosmedReport:
    def __init__(self,
                  patient_name="Шмидхубер Юрген Иванович",
                  doctor_name="Пупкин В.П.",
                  date=str(datetime.now()),
                  description="",
                  pathologies=[]):
        self.original_UID =  "1.2.840.10008.5.1.4.1.1.88.33"
        self.UID = generate_UID(original_UID=self.original_UID, add_id="2")
        modification_time = time.strftime("%H%M%S")
        modification_date = time.strftime("%Y%m%d")
        ds = Dataset()
        ds.SeriesInstanceUID = self.UID
        ds.SOPClassUID = "1.2.840.10008.5.1.4.1.1.88.33"
        ds.SOPInstanceUID = self.UID
        ds.StationName = "Mosmed"
        ds.StudyDate = get_formatted_string(modification_date)
        ds.ContentDate = get_formatted_string(modification_date)
        ds.StudyTime = get_formatted_string(date)
        ds.ContentTime = get_formatted_string(date)
        ds.StudyInstanceUID = generate_UID(self.original_UID,add_id="1")
        ds.StudyID = "404"
        ds.PatientName = get_formatted_string(patient_name)
        ds.PatientSex = get_formatted_string("M")

        ds.PatientBirthDate= get_formatted_string("09.03.1997")

        # HEADER BEGIN
        ds.Modality = "SR"
        ds.StudyDate = modification_date  # non formatted study data
        # HEADER END
        content_sequence = generate_sequence_report(patient_name,doctor_name,date,description,pathologies)
        ds.ContentSequence = content_sequence
        # ds.file_meta = file_meta
        ds.file_meta = create_file_meta(ds.SOPInstanceUID)
        ds.is_implicit_VR = True
        ds.is_little_endian = True
        ds.SpecificCharacterSet = 'ISO_IR 192'
        self.ds = ds

    def save(self, output_dir, name=None):
        if not name:
            name = self.ds.SOPInstanceUID
        # Write to the output directory and add the extension dcm, to force writing in DICOM format.
        dicom_filename = "{0}.dcm".format(str(name))
        dicom_path = os.path.join(output_dir, dicom_filename)
        self.ds.save_as(dicom_path, write_like_original=False)
        # fix_preamble(dicom_path)

        return dicom_path


def fix_preamble(dicom_path):
    fp = BytesIO()
    fp.write(b'\x00' * 128)
    fp.write(b'DICM')

    # Add the contents of the file
    f = open(dicom_path, 'rb')
    fp.write(f.read())
    f.close()
    fp.seek(0)

    with open(dicom_path, 'wb') as out:
        out.write(fp.read())


def generate_sequence_report(patient_name="Шмидхубер Юрген Иванович",
                  doctor_name="Пупкин В.П.",
                  date=str(datetime.now()),
                  description="",
                  pathologies=[]):
    if len(pathologies)==0:
        res = Content(CodeMeaning='Заключение', TextValue="""Рентгенологические признаки очаговой патологии органов грудной клетки не выявлены.""")
    else:
        res = Content(CodeMeaning='Заключение', TextValue="""Рентгенологические признаки очаговой патологии органов грудной клетки обнаружены, необходима консультация специалиста.""")
    return [
        Content(CodeMeaning=f'ФИО: {patient_name}', TextValue="""Дата рождения: 09.03.1997г.;
        Пол: М 0,04 мЗв РГ органов грудной клетки в прямой проекции в ортопозиции."""),
        Content(CodeMeaning='Описание', TextValue=description),
        res,
        Content(CodeMeaning="""_______________________________________________________________\r""",
                TextValue=f"""Врач: {doctor_name} Дата: {date} 8:31 Подпись:__________"""),
    ]


def create_file_meta(MediaStorageSOPInstanceUID):
    file_meta = Dataset()
    file_meta.FileMetaInformationGroupLength = 174
    file_meta.FileMetaInformationVersion = b'\x00\x01'
    file_meta.MediaStorageSOPClassUID = "1.2.840.10008.5.1.4.1.1.88.33"
    file_meta.MediaStorageSOPInstanceUID = MediaStorageSOPInstanceUID
    file_meta.ImplementationClassUID = "1.2.40.0.13.1.1"
    file_meta.TransferSyntaxUID = "1.2.840.10008.1.2"

    return file_meta


def get_report_text(disease_value, target_disease, THRESHOLD):
    disease_found = disease_value > THRESHOLD
    return reports["Pathology"][disease_found].format(str(disease_value))


class Content(Dataset):
    def __init__(self, RelationshipType='CONTAINS', ValueType='TEXT',
                 CodeValue='209001', CodingSchemeDesignator='99PMP', CodeMeaning='Название сервиса',
                 TextValue='AI Example'):
        super().__init__()
        self.RelationshipType = RelationshipType
        self.ValueType = ValueType
        self.ConceptNameCodeSequence = [ConceptNameCodeSequence(CodeValue, CodingSchemeDesignator, CodeMeaning)]
        self.TextValue = TextValue


class ConceptNameCodeSequence(Dataset):
    def __init__(self, CodeValue, CodingSchemeDesignator, CodeMeaning):
        super().__init__()
        self.CodeValue = CodeValue
        self.CodingSchemeDesignator = CodingSchemeDesignator
        self.CodeMeaning = CodeMeaning


def current_time():
    return datetime.datetime.now().astimezone().isoformat(timespec="milliseconds")[:-3] + "00"


def current_date_time():
    return datetime.datetime.now().astimezone(pytz.timezone('Europe/Moscow')).strftime("%Y-%m-%d %H:%M")   